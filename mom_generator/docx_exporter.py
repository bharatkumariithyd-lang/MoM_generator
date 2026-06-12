"""
docx_exporter.py
================

Step 3 of the pipeline: take the structured MoM dictionary and write a nicely
formatted Word (.docx) document using python-docx.

The layout we produce:

    MINUTES OF MEETING            (centered title)
    <Meeting Title>

    Date & Time:  ...
    Attendees:    ...

    1. Agenda
    2. Key Discussion Points
    3. Decisions Made
    4. Action Items            (as a table: Task | Owner | Deadline)
    5. Next Meeting
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# A calm, professional dark-blue for headings.
_HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)


def _add_bullet_list(document, items):
    """Add a bulleted list. If empty, add a gentle placeholder line."""
    if not items:
        document.add_paragraph("Not mentioned in the meeting.", style="List Bullet")
        return
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _add_section_heading(document, text):
    """Add a styled section heading and return it (so callers can tweak it)."""
    heading = document.add_heading(text, level=1)
    # Recolor the heading run(s) to our brand color for a polished look.
    for run in heading.runs:
        run.font.color.rgb = _HEADING_COLOR
    return heading


def export_to_docx(mom: dict, output_path: str) -> str:
    """
    Write the Minutes of Meeting dictionary to a .docx file.

    Parameters
    ----------
    mom : dict
        The structured minutes from mom_builder.build_mom().
    output_path : str
        Where to save the .docx file.

    Returns
    -------
    str  (the output_path, for convenience)
    """
    document = Document()

    # ---- Title block ------------------------------------------------------
    title = document.add_heading("MINUTES OF MEETING", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(mom.get("meeting_title", "Meeting Minutes"))
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)

    # ---- Meta block: date/time and attendees ------------------------------
    # We use a small helper to print "Label: value", leaving a blank for the
    # user to fill in when the audio did not mention it.
    date_time = mom.get("date_time", "") or "______________________  (fill in)"
    meta = document.add_paragraph()
    meta.add_run("Date & Time: ").bold = True
    meta.add_run(date_time)

    attendees = mom.get("attendees", [])
    attendees_text = ", ".join(attendees) if attendees else "______________________  (fill in)"
    meta2 = document.add_paragraph()
    meta2.add_run("Attendees: ").bold = True
    meta2.add_run(attendees_text)

    # ---- Numbered content sections ---------------------------------------
    _add_section_heading(document, "Agenda")
    _add_bullet_list(document, mom.get("agenda_items", []))

    _add_section_heading(document, "Key Discussion Points")
    _add_bullet_list(document, mom.get("key_discussion_points", []))

    _add_section_heading(document, "Decisions Made")
    _add_bullet_list(document, mom.get("decisions_made", []))

    # ---- Action Items as a table -----------------------------------------
    _add_section_heading(document, "Action Items")
    action_items = mom.get("action_items", [])
    if action_items:
        table = document.add_table(rows=1, cols=3)
        table.style = "Light List Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Task"
        header_cells[1].text = "Owner"
        header_cells[2].text = "Deadline"
        for cell in header_cells:
            # Make the header text bold.
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # One row per action item.
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("task", ""))
            row_cells[1].text = str(item.get("owner", "")) or "(unassigned)"
            row_cells[2].text = str(item.get("deadline", "")) or "(no deadline)"
    else:
        document.add_paragraph("No action items were recorded.", style="List Bullet")

    # ---- Next meeting -----------------------------------------------------
    _add_section_heading(document, "Next Meeting")
    next_meeting = mom.get("next_meeting", "")
    document.add_paragraph(next_meeting if next_meeting else "Not mentioned in the meeting.")

    # ---- Footer note ------------------------------------------------------
    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "\nGenerated automatically from the meeting audio. "
        "Please review and edit before sharing."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    document.save(output_path)
    print(f"[docx_exporter] Minutes saved to: {output_path}")
    return output_path
